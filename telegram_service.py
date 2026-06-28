"""
Telegram Bot Service — v4 with timeout preference onboarding step.
Onboarding now has 5 steps:
  1. repo (owner/repo)
  2. branch
  3. github token
  4. timeout duration (hours)
  5. timeout action (accept / decline)
  hi there is there anyone there is there ai checking on this
"""

import asyncio
import json
import os
import re
import tempfile
from datetime import datetime, timezone
from typing import Any, Callable, Coroutine, Dict, Optional

import httpx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ai_service import CommitDecision, ai_service
from config import CONFIG, logger
import database as db


class TelegramServiceError(Exception):
    pass

class TelegramAPIError(TelegramServiceError):
    pass


class TelegramService:

    EMOJI = {
        "pending":       "⏳",
        "accept":        "✅",
        "decline":       "❌",
        "report":        "📊",
        "risk_critical": "🔴",
        "risk_high":     "🟠",
        "risk_medium":   "🟡",
        "risk_low":      "🟢",
        "loading":       "🔄",
        "warning":       "⚠️",
        "rollback":      "⏪",
        "commit":        "📝",
        "user":          "👤",
        "clock":         "🕐",
        "files":         "📁",
        "ai":            "🤖",
        "rocket":        "🚀",
        "key":           "🔑",
        "link":          "🔗",
        "check":         "☑️",
        "timeout":       "⏰",
    }

    # ── Persistent reply keyboard shown after setup and via /menu ─────────────
    MAIN_MENU_KEYBOARD = {
        "keyboard": [
            [{"text": "👤 My Profile"}, {"text": "📜 Commit History"}],
            [{"text": "📊 Active Reviews"}, {"text": "⚙️ Settings"}],
            [{"text": "🔍 Full Code Analysis"}, {"text": "👥 Author Performance"}],
            [{"text": "📞 Contact Support"}, {"text": "🙈 Hide Menu"}],
        ],
        "resize_keyboard": True,
        "is_persistent": True,
    }

    # Sent to dismiss the reply keyboard; user can get it back with /menu.
    HIDE_KEYBOARD = {"remove_keyboard": True}

    def __init__(self) -> None:
        self.token    = CONFIG.telegram_bot_token
        self.base_url = f"https://api.telegram.org/bot{self.token}"
        self._client: Optional[httpx.AsyncClient] = None

    # ── HTTP ──────────────────────────────────────────────────────────────────

    async def _get_client(self) -> httpx.AsyncClient:
        if self._client is None or self._client.is_closed:
            self._client = httpx.AsyncClient(timeout=30.0)
        return self._client

    async def close(self) -> None:
        if self._client and not self._client.is_closed:
            await self._client.aclose()
            self._client = None

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _safe_error(self, exc: Exception) -> str:
        """
        Fix #12: strip anything that looks like a GitHub token from exception
        messages before they are sent to the user's Telegram chat.
        """
        msg = str(exc)
        msg = re.sub(r"(ghp_|github_pat_)[A-Za-z0-9_]+", "[REDACTED]", msg)
        msg = re.sub(r"token [A-Za-z0-9_]{10,}", "token [REDACTED]", msg)
        return msg[:300]

    def _risk_emoji(self, risk_level: str) -> str:
        return {
            "critical": self.EMOJI["risk_critical"],
            "high":     self.EMOJI["risk_high"],
            "medium":   self.EMOJI["risk_medium"],
            "low":      self.EMOJI["risk_low"],
        }.get(risk_level.lower(), self.EMOJI["warning"])

    def _is_revert_commit(self, commit_metadata: Dict[str, Any]) -> bool:
        return commit_metadata.get("message", "").startswith("revert: rollback commit ")

    def _extract_reverted_sha(self, commit_metadata: Dict[str, Any]) -> Optional[str]:
        m = re.search(r"revert: rollback commit ([0-9a-f]{7,40})", commit_metadata.get("message", ""))
        return m.group(1) if m else None

    # ── Messaging primitives ──────────────────────────────────────────────────

    async def send_message(
        self,
        chat_id: str,
        text: str,
        reply_markup: Optional[Dict] = None,
        parse_mode: str = "Markdown",
    ) -> int:
        payload: Dict[str, Any] = {
            "chat_id":                  chat_id,
            "text":                     text,
            "parse_mode":               parse_mode,
            "disable_web_page_preview": True,
        }
        if reply_markup:
            payload["reply_markup"] = json.dumps(reply_markup)
        try:
            client   = await self._get_client()
            response = await client.post(f"{self.base_url}/sendMessage", json=payload)
            response.raise_for_status()
            result = response.json()
            if not result.get("ok"):
                raise TelegramAPIError(f"Telegram error: {result.get('description')}")
            return result["result"]["message_id"]
        except httpx.HTTPError as exc:
            raise TelegramAPIError(f"Failed to send message: {exc}") from exc

    async def edit_message(
        self,
        chat_id: str,
        message_id: int,
        text: str,
        reply_markup: Optional[Dict] = None,
    ) -> None:
        payload: Dict[str, Any] = {
            "chat_id":                  chat_id,
            "message_id":               message_id,
            "text":                     text,
            "parse_mode":               "Markdown",
            "disable_web_page_preview": True,
        }
        if reply_markup is not None:
            payload["reply_markup"] = json.dumps(reply_markup)
        try:
            client = await self._get_client()
            r = await client.post(f"{self.base_url}/editMessageText", json=payload)
            # Fix #23: if the original card was deleted by the user, fall back to
            # sending a fresh message so the outcome is never silently lost.
            if r.status_code == 400:
                body = r.json()
                desc = body.get("description", "").lower()
                if "message to edit not found" in desc or "message_id_invalid" in desc:
                    logger.warning(
                        "Message %d not found for chat %s — sending fallback message",
                        message_id, chat_id,
                    )
                    await self.send_message(chat_id, text, reply_markup=reply_markup)
        except httpx.HTTPError as exc:
            logger.warning("Could not edit message %d: %s", message_id, exc)

    async def delete_message(self, chat_id: str, message_id: int) -> None:
        try:
            client = await self._get_client()
            await client.post(
                f"{self.base_url}/deleteMessage",
                json={"chat_id": chat_id, "message_id": message_id},
            )
        except httpx.HTTPError as exc:
            logger.warning("Could not delete message %d: %s", message_id, exc)

    async def send_processing(self, chat_id: str, text: str) -> int:
        return await self.send_message(chat_id, text)

    async def answer_callback(
        self, callback_query_id: str, text: Optional[str] = None, show_alert: bool = False
    ) -> None:
        payload: Dict[str, Any] = {"callback_query_id": callback_query_id}
        if text:
            payload["text"] = text
        if show_alert:
            payload["show_alert"] = True
        try:
            client = await self._get_client()
            await client.post(f"{self.base_url}/answerCallbackQuery", json=payload)
        except httpx.HTTPError as exc:
            logger.warning("Could not answer callback: %s", exc)

    async def send_document(
        self, chat_id: str, file_path: str,
        caption: Optional[str] = None, filename: Optional[str] = None,
    ) -> None:
        fname = filename or os.path.basename(file_path)
        try:
            client = await self._get_client()
            with open(file_path, "rb") as fh:
                files = {"document": (fname, fh, "application/octet-stream")}
                data: Dict[str, Any] = {"chat_id": chat_id}
                if caption:
                    data["caption"] = caption[:1024]
                response = await client.post(f"{self.base_url}/sendDocument", data=data, files=files)
            response.raise_for_status()
        except httpx.HTTPError as exc:
            raise TelegramAPIError(f"Failed to send document: {exc}") from exc

    # ── Main menu ─────────────────────────────────────────────────────────────

    async def show_main_menu(self, chat_id: str, text: str = "Choose an option:") -> None:
        """Send (or re-send) the persistent reply keyboard."""
        await self.send_message(chat_id, text, reply_markup=self.MAIN_MENU_KEYBOARD)

    async def hide_main_menu(self, chat_id: str) -> None:
        """Dismiss the reply keyboard. User can restore it anytime with /menu."""
        await self.send_message(
            chat_id,
            "Menu hidden. Send /menu whenever you want it back.",
            reply_markup=self.HIDE_KEYBOARD,
        )

    # ── Onboarding wizard (5 steps) ───────────────────────────────────────────

    async def handle_cancel(self, chat_id: str) -> None:
        """Fix #22: escape hatch — lets user abort onboarding at any step."""
        user = db.get_user(chat_id)
        if user and user.get("onboard_step") not in (None, "done"):
            db.upsert_user(chat_id, onboard_step="done")
            await self.send_message(chat_id, "Setup cancelled. Send /start to begin again.")
        else:
            await self.send_message(chat_id, "Nothing to cancel. Send /start to configure.")

    async def handle_reconnect(self, chat_id: str) -> None:
        """Refresh an expired/revoked token without wiping repo/branch/timeout settings."""
        user = db.get_user(chat_id)
        if not user or not user.get("owner"):
            await self.send_message(chat_id, "You're not set up yet. Send /start to begin.")
            return
        db.upsert_user(chat_id, onboard_step="await_reconnect_token")
        await self.send_message(
            chat_id,
            f"{self.EMOJI['key']} *Reconnect GitHub*\n\n"
            f"Currently monitoring `{user['owner']}/{user['repo']}` ({user.get('branch','main')}) — "
            "this won't change.\n\n"
            "📋 *How to create a new token:*\n"
            "1. GitHub → avatar (top-right) → *Settings*\n"
            "2. Left sidebar → *Developer settings*\n"
            "3. *Personal access tokens* → *Tokens (classic)*\n"
            "4. *Generate new token (classic)*\n"
            "5. Name it `CommitGuardian`, set expiry (90 days)\n"
            "6. Check these boxes:\n"
            "   ☑️ `repo` — full repo access\n"
            "   ☑️ `read:user` — read your username\n"
            "7. *Generate token* — copy it immediately!\n\n"
            "Paste your new token:",
        )

    async def handle_start(self, chat_id: str) -> None:
        # Fix #9: warn the user if they have pending reviews before wiping config.
        user = db.get_user(chat_id)
        if user and db.get_active_reviews_for_user(chat_id):
            db.upsert_user(chat_id, onboard_step="await_restart_confirm")
            await self.send_message(
                chat_id,
                f"{self.EMOJI['warning']} *You have pending reviews.*\n\n"
                "Reconfiguring now will clear your GitHub token and break any "
                "auto-rollbacks for those reviews.\n\n"
                "Send `CONFIRM` to proceed, or /status to review them first.",
            )
            return

        await self._do_start(chat_id)

    async def _do_start(self, chat_id: str) -> None:
        """Actually reset config and start the onboarding wizard."""
        db.upsert_user(
            chat_id,
            onboard_step="await_repo",
            owner=db.CLEAR, repo=db.CLEAR, branch="main",
            github_token=db.CLEAR,
            timeout_hours=24,
            timeout_action="accept",
        )
        db.clear_token_alert(chat_id)  # reset cooldown when user fully reconfigures
        await self.send_message(
            chat_id,
            f"{self.EMOJI['rocket']} *Welcome to Commit Guardian!*\n\n"
            "I review every GitHub commit with AI and send it here for your approval.\n\n"
            f"{self.EMOJI['commit']} *Step 1 of 5 — Repository*\n\n"
            "Send your GitHub repo:\n`owner/repository-name`\n\n"
            "_Example: `torvalds/linux`_",
        )

    async def handle_message(self, chat_id: str, text: str) -> bool:
        """Returns True if message was consumed by the wizard."""
        user = db.get_user(chat_id)
        if not user:
            return False

        step = user.get("onboard_step", "done")
        if step == "done":
            return False

        # Fix #22: /cancel is an escape hatch at any onboarding step.
        if text.strip().lower() in ("/cancel", "/stop", "cancel"):
            await self.handle_cancel(chat_id)
            return True

        # Fix #9: user confirmed they want to reconfigure despite pending reviews.
        if step == "await_restart_confirm":
            if text.strip().upper() == "CONFIRM":
                await self._do_start(chat_id)
            else:
                await self.send_message(
                    chat_id,
                    f"{self.EMOJI['warning']} Send `CONFIRM` to proceed with reconfiguration, "
                    "or /status to see your pending reviews.",
                )
            return True

        # ── Step 1: repo ──────────────────────────────────────────────────────
        if step == "await_repo":
            if "/" not in text or text.count("/") != 1:
                await self.send_message(
                    chat_id,
                    f"{self.EMOJI['warning']} Please use the format `owner/repo`.\n"
                    "_Example: `torvalds/linux`_",
                )
                return True
            owner, repo = text.strip().split("/", 1)
            owner = owner.strip()
            repo  = repo.strip()
            # Fix #14: enforce GitHub's allowed character set so malformed values
            # are never stored in the DB.
            if not re.match(r"^[a-zA-Z0-9._-]{1,100}$", owner) or not re.match(r"^[a-zA-Z0-9._-]{1,100}$", repo):
                await self.send_message(
                    chat_id,
                    f"{self.EMOJI['warning']} Invalid format — use only letters, numbers, "
                    "hyphens, dots, or underscores (max 100 chars each).\nTry again:",
                )
                return True
            db.upsert_user(chat_id, owner=owner, repo=repo, onboard_step="await_branch")
            await self.send_message(
                chat_id,
                f"{self.EMOJI['check']} Repo: `{owner}/{repo}`\n\n"
                f"{self.EMOJI['commit']} *Step 2 of 5 — Branch*\n\n"
                "Which branch to monitor?\n_(Type name or send `main`)_",
            )
            return True

        # ── Step 2: branch ────────────────────────────────────────────────────
        if step == "await_branch":
            branch = text.strip() or "main"
            db.upsert_user(chat_id, branch=branch, onboard_step="await_token")
            await self.send_message(
                chat_id,
                f"{self.EMOJI['check']} Branch: `{branch}`\n\n"
                f"{self.EMOJI['key']} *Step 3 of 5 — GitHub Personal Access Token*\n\n"
                "I need a token to read commits and perform rollbacks.\n\n"
                "📋 *How to create one:*\n"
                "1. GitHub → avatar (top-right) → *Settings*\n"
                "2. Left sidebar → *Developer settings*\n"
                "3. *Personal access tokens* → *Tokens (classic)*\n"
                "4. *Generate new token (classic)*\n"
                "5. Name it `CommitGuardian`, set expiry (90 days)\n"
                "6. Check these boxes:\n"
                "   ☑️ `repo` — full repo access\n"
                "   ☑️ `read:user` — read your username\n"
                "7. *Generate token* — copy it immediately!\n\n"
                "⚠️ _GitHub only shows it once._\n\n"
                "Paste your token:",
            )
            return True

        # ── Reconnect: replace just the token, keep everything else ───────────
        if step == "await_reconnect_token":
            token = text.strip()
            if not (token.startswith("ghp_") or token.startswith("github_pat_")):
                await self.send_message(
                    chat_id,
                    f"{self.EMOJI['warning']} Doesn't look right "
                    "(should start with `ghp_` or `github_pat_`).\nTry again:",
                )
                return True
            valid, username = await self._validate_github_token(token)
            if not valid:
                await self.send_message(
                    chat_id,
                    f"{self.EMOJI['warning']} Token invalid or expired — GitHub returned an error.\n"
                    "Please generate a new one and paste it here:",
                )
                return True
            db.upsert_user(chat_id, github_token=token, onboard_step="done")
            db.clear_token_alert(chat_id)   # allow fresh alert if token breaks again
            greeting = f" (logged in as `{username}`)" if username else ""
            await self.send_message(
                chat_id,
                f"{self.EMOJI['check']} Token refreshed{greeting}. "
                "I'll keep using your existing repo, branch, and timeout settings.\n\n"
                "You're back up and running 🚀",
            )
            return True

        # ── Step 3: token ─────────────────────────────────────────────────────
        if step == "await_token":
            token = text.strip()
            if not (token.startswith("ghp_") or token.startswith("github_pat_")):
                await self.send_message(
                    chat_id,
                    f"{self.EMOJI['warning']} Doesn't look right "
                    "(should start with `ghp_` or `github_pat_`).\nTry again:",
                )
                return True
            # Fix #16: validate token immediately with a live GitHub API call so
            # truncated / expired / revoked tokens are caught during onboarding.
            valid, username = await self._validate_github_token(token)
            if not valid:
                await self.send_message(
                    chat_id,
                    f"{self.EMOJI['warning']} Token invalid or expired — GitHub returned an error.\n"
                    "Please generate a new one and paste it here:",
                )
                return True
            db.upsert_user(chat_id, github_token=token, onboard_step="await_timeout_hours")
            greeting = f" (logged in as `{username}`)" if username else ""
            await self.send_message(
                chat_id,
                f"{self.EMOJI['check']} Token verified{greeting}.\n\n"
                f"{self.EMOJI['timeout']} *Step 4 of 5 — Review Timeout*\n\n"
                "If you don't respond to a review, how many hours before I auto-act?\n\n"
                "Common choices:\n"
                "• `6` — half a day\n"
                "• `24` — one day _(recommended)_\n"
                "• `48` — two days\n"
                "• `0` — never auto-act\n\n"
                "Send a number:",
            )
            return True

        # ── Step 4: timeout hours ─────────────────────────────────────────────
        if step == "await_timeout_hours":
            try:
                hours = int(text.strip())
                if hours < 0:
                    raise ValueError
            except ValueError:
                await self.send_message(
                    chat_id,
                    f"{self.EMOJI['warning']} Please send a whole number (e.g. `24`). "
                    "Send `0` to disable auto-action.",
                )
                return True

            db.upsert_user(chat_id, timeout_hours=hours, onboard_step="await_timeout_action")

            if hours == 0:
                # Skip action step — auto-action is disabled
                db.upsert_user(chat_id, timeout_action="none", onboard_step="done")
                await self._finish_onboarding(chat_id)
                return True

            await self.send_message(
                chat_id,
                f"{self.EMOJI['check']} Timeout: `{hours}` hours\n\n"
                f"{self.EMOJI['timeout']} *Step 5 of 5 — Timeout Action*\n\n"
                f"After `{hours}` hours with no response, what should I do?\n\n"
                "• Send `accept` — keep the commit in the repo _(safer, won't disrupt teammates)_\n"
                "• Send `decline` — automatically roll it back _(stricter, security-first)_",
            )
            return True

        # ── Step 5: timeout action ────────────────────────────────────────────
        if step == "await_timeout_action":
            action = text.strip().lower()
            if action not in ("accept", "decline"):
                await self.send_message(
                    chat_id,
                    f"{self.EMOJI['warning']} Please send exactly `accept` or `decline`.",
                )
                return True

            db.upsert_user(chat_id, timeout_action=action, onboard_step="done")
            await self._finish_onboarding(chat_id)
            return True

        # ── Branch-only update (from My Profile edit) ─────────────────────────
        if step == "await_branch_update":
            branch = text.strip() or "main"
            db.upsert_user(chat_id, branch=branch, onboard_step="done")
            await self.send_message(chat_id, f"{self.EMOJI['check']} Branch updated to `{branch}`.")
            await self.handle_my_profile(chat_id)
            return True

        return False

    async def _validate_github_token(self, token: str) -> tuple:
        """
        Fix #16: call the GitHub /user endpoint to verify the token is live.
        Returns (True, login_username) or (False, "").
        """
        try:
            async with httpx.AsyncClient(timeout=15.0) as client:
                r = await client.get(
                    "https://api.github.com/user",
                    headers={"Authorization": f"token {token}", "User-Agent": "CommitGuardian/2.0"},
                )
                if r.status_code == 200:
                    return True, r.json().get("login", "")
                return False, ""
        except Exception as exc:
            logger.warning("GitHub token validation request failed: %s", exc)
            return False, ""

    async def _finish_onboarding(self, chat_id: str) -> None:
        user           = db.get_user(chat_id)
        webhook_secret = user["webhook_secret"]
        owner          = user["owner"]
        repo           = user["repo"]
        branch         = user["branch"]
        timeout_hours  = user["timeout_hours"]
        timeout_action = user["timeout_action"]
        webhook_url    = f"{CONFIG.public_url}/webhook/github/{chat_id}"

        if timeout_action == "none" or timeout_hours == 0:
            timeout_summary = "Disabled — reviews never auto-expire"
        else:
            timeout_summary = f"Auto-*{timeout_action}* after `{timeout_hours}` hours"

        await self.send_message(
            chat_id,
            f"{self.EMOJI['rocket']} *You're all set!*\n\n"
            f"*Summary:*\n"
            f"• Repo: `{owner}/{repo}`\n"
            f"• Branch: `{branch}`\n"
            f"• Timeout: {timeout_summary}\n\n"
            f"───────────────────────────\n"
            f"{self.EMOJI['link']} *Add the webhook to GitHub:*\n\n"
            f"1. Go to `{owner}/{repo}` → *Settings* → *Webhooks* → *Add webhook*\n"
            f"2. Fill in:\n"
            f"   • *Payload URL:* `{webhook_url}`\n"
            f"   • *Content type:* `application/json`\n"
            f"   • *Secret:* `{webhook_secret}`\n"
            f"   • *Which events?* → *Just the push event*\n"
            f"   • *Active* ✅\n"
            f"3. Click *Add webhook*\n\n"
            f"⚠️ _Copy the secret above now and delete this message after saving it._\n\nEvery push to `{branch}` will now appear here for review. 🎉\n\n"
            f"_Commands: /status · /settings · /start (reconfigure)_",
        )
        await self.show_main_menu(chat_id, "✅ Setup complete — use the menu below:")

    # ── /settings command ─────────────────────────────────────────────────────

    async def handle_settings(self, chat_id: str) -> None:
        """Show current settings with quick-change options."""
        user = db.get_user(chat_id)
        if not user or not db.is_setup_complete(chat_id):
            await self.send_message(chat_id, "You're not set up yet. Send /start to begin.")
            return

        timeout_action = user.get("timeout_action", "accept")
        timeout_hours  = user.get("timeout_hours", 24)

        if timeout_action == "none" or timeout_hours == 0:
            timeout_str = "Disabled"
        else:
            timeout_str = f"Auto-{timeout_action} after {timeout_hours}h"

        keyboard = {
            "inline_keyboard": [
                [
                    {"text": "⏰ Change timeout",        "callback_data": "cfg:timeout"},
                    {"text": "🔄 Full reconfigure",      "callback_data": "cfg:restart"},
                ],
            ]
        }

        await self.send_message(
            chat_id,
            f"⚙️ *Your Settings*\n\n"
            f"• Repo: `{user.get('owner')}/{user.get('repo')}`\n"
            f"• Branch: `{user.get('branch','main')}`\n"
            f"• Timeout: {timeout_str}\n\n"
            f"Webhook URL:\n`{CONFIG.public_url}/webhook/github/{chat_id}`",
            reply_markup=keyboard,
        )

    # ── My Profile ────────────────────────────────────────────────────────────

    async def handle_my_profile(self, chat_id: str) -> None:
        """Show full profile with masked token and inline edit buttons."""
        user = db.get_user(chat_id)
        if not user or not db.is_setup_complete(chat_id):
            await self.send_message(
                chat_id,
                "👤 *My Profile*\n\n"
                "You haven't completed setup yet. Send /start to configure your repo.",
            )
            return

        token = user.get("github_token") or ""
        if len(token) > 12:
            masked_token = f"{token[:8]}...{token[-4:]}"
        elif token:
            masked_token = "****"
        else:
            masked_token = "_not set_"

        timeout_hours  = user.get("timeout_hours", 24)
        timeout_action = user.get("timeout_action", "accept")
        if timeout_hours == 0 or timeout_action == "none":
            timeout_str = "Disabled"
        else:
            timeout_str = f"Auto-{timeout_action} after {timeout_hours}h"

        webhook_url = f"{CONFIG.public_url}/webhook/github/{chat_id}"
        created_at  = (user.get("created_at") or "")[:10] or "N/A"

        keyboard = {
            "inline_keyboard": [
                [
                    {"text": "🔑 Refresh Token",    "callback_data": "prof:token"},
                    {"text": "🔀 Change Branch",    "callback_data": "prof:branch"},
                ],
                [
                    {"text": "⏰ Change Timeout",   "callback_data": "prof:timeout"},
                    {"text": "🔄 Full Reconfigure", "callback_data": "prof:restart"},
                ],
            ]
        }

        await self.send_message(
            chat_id,
            f"👤 *My Profile*\n\n"
            f"• {self.EMOJI['commit']} *Repository:* `{user.get('owner')}/{user.get('repo')}`\n"
            f"• 🌿 *Branch:* `{user.get('branch', 'main')}`\n"
            f"• {self.EMOJI['key']} *GitHub Token:* `{masked_token}`\n"
            f"• {self.EMOJI['timeout']} *Timeout:* {timeout_str}\n"
            f"• 📅 *Member since:* {created_at}\n\n"
            f"🔗 *Webhook URL:*\n`{webhook_url}`\n\n"
            f"_Use the buttons below to modify your settings:_",
            reply_markup=keyboard,
        )

    # ── Commit History ────────────────────────────────────────────────────────

    async def handle_commit_history(self, chat_id: str) -> None:
        """Show paginated commit history with AI decisions."""
        user = db.get_user(chat_id)
        if not user or not db.is_setup_complete(chat_id):
            await self.send_message(
                chat_id,
                "📜 *Commit History*\n\n"
                "You haven't completed setup yet. Send /start to configure your repo.",
            )
            return

        reviews = db.get_all_reviews_for_user(chat_id, limit=15)

        if not reviews:
            await self.send_message(
                chat_id,
                "📜 *Commit History*\n\n"
                "_No commits reviewed yet._\n\n"
                "Once you push to your repository, commit reviews will appear here.",
            )
            return

        status_icon = {
            "accepted": "✅",
            "declined": "❌",
            "pending":  "⏳",
        }

        lines: list = []
        for i, r in enumerate(reviews, 1):
            sha    = r["commit_sha"][:7]
            status = r.get("status", "pending")
            icon   = status_icon.get(status, "⏳")
            date   = (r.get("created_at") or "")[:10]

            # Pull commit message from stored metadata
            try:
                meta = json.loads(r.get("commit_meta_json") or "{}")
            except Exception:
                meta = {}
            commit_msg = (meta.get("message") or "")[:55]

            # Pull risk from stored decision
            try:
                dec = json.loads(r.get("decision_json") or "{}")
            except Exception:
                dec = {}
            risk       = (dec.get("risk_level") or "").upper()
            risk_emoji = self._risk_emoji(risk.lower()) if risk else ""

            line = f"{i}. {icon} `{sha}` {risk_emoji} _{date}_ — *{status.capitalize()}*"
            if commit_msg:
                # Escape backtick to avoid markdown breakage
                safe_msg = commit_msg.replace("`", "'")
                line += f"\n    `{safe_msg}`"
            lines.append(line)

        header = f"📜 *Commit History* (last {len(reviews)})\n"
        body   = "\n\n".join(lines)
        text   = f"{header}Repo: `{user.get('owner')}/{user.get('repo')}`\n\n{body}"

        if len(text) > 4000:
            text = text[:4000] + "\n\n_…truncated_"

        await self.send_message(chat_id, text)

    # ── Contact Support ───────────────────────────────────────────────────────

    async def handle_contact_support(self, chat_id: str) -> None:
        """Send a message with a link to the support Telegram account."""
        keyboard = {
            "inline_keyboard": [
                [{"text": "💬 Open Support Chat", "url": "https://t.me/mikoz_124"}],
            ]
        }
        await self.send_message(
            chat_id,
            "📞 *Contact Support*\n\n"
            "Having trouble or need help with Commit Guardian?\n\n"
            "Reach out to our support team directly on Telegram.\n\n"
            "_Tap the button below to open a chat with us._",
            reply_markup=keyboard,
        )

    # ── Review request ────────────────────────────────────────────────────────

    async def send_review_request(
        self,
        chat_id: str,
        commit_metadata: Dict[str, Any],
        decision: CommitDecision,
    ) -> int:
        """Returns message_id or -1 if skipped."""
        if self._is_revert_commit(commit_metadata):
            reverted = self._extract_reverted_sha(commit_metadata)
            await self.send_message(
                chat_id,
                f"{self.EMOJI['rollback']} *Rollback confirmed*\n\n"
                f"Revert commit `{commit_metadata.get('sha','')[:7]}` applied.\n"
                f"_This reverted commit `{reverted}`. No review needed._",
            )
            return -1

        user          = db.get_user(chat_id) or {}
        timeout_hours = user.get("timeout_hours", 24)
        timeout_action = user.get("timeout_action", "accept")
        sha           = commit_metadata.get("sha", "unknown")
        files         = commit_metadata.get("files", [])
        stats         = commit_metadata.get("stats", {})

        files_summary = "\n".join(
            f"  `{f['filename']}` ({f['status']}, +{f['additions']}/-{f['deletions']})"
            for f in files[:15]
        )
        if len(files) > 15:
            files_summary += f"\n  … and {len(files) - 15} more"

        concerns  = "\n".join(f"  • {c}" for c in decision.concerns[:5])       or "  None"
        positives = "\n".join(f"  • {p}" for p in decision.positive_aspects[:5]) or "  None"

        if timeout_hours == 0 or timeout_action == "none":
            timeout_notice = "_No auto-action — this review won't expire._"
        else:
            timeout_notice = (
                f"_{self.EMOJI['timeout']} No response in *{timeout_hours}h* → "
                f"will auto-*{timeout_action}*._"
            )

        text = (
            f"{self.EMOJI['commit']} *New Commit Requires Review*\n\n"
            f"{self.EMOJI['user']} *Author:* {commit_metadata.get('author_name','Unknown')}"
            f" (`{commit_metadata.get('author_email','N/A')}`)\n"
            f"{self.EMOJI['clock']} *Committed:* {commit_metadata.get('committed_at','N/A')}\n"
            f"{self.EMOJI['commit']} *SHA:* `{sha[:7]}`\n"
            f"{self.EMOJI['files']} *Files:* {len(files)} changed  "
            f"(+{stats.get('additions',0)} / -{stats.get('deletions',0)} lines)\n\n"
            f"💬 *Message:*\n```\n{commit_metadata.get('message','')[:300]}\n```\n\n"
            f"———\n\n"
            f"{self.EMOJI['ai']} *AI Assessment*\n\n"
            f"{self._risk_emoji(decision.risk_level)} *Risk:* {decision.risk_level.upper()}\n"
            f"🎯 *Confidence:* {decision.confidence_score:.0%}\n"
            f"📋 *Suggestion:* {decision.suggested_action}\n\n"
            f"📝 *Summary:* {decision.summary}\n\n"
            f"⚠️ *Concerns:*\n{concerns}\n\n"
            f"✨ *Positives:*\n{positives}\n\n"
            f"———\n*Changed Files:*\n{files_summary}\n\n"
            f"{timeout_notice}\n\n"
            f"_Select an action:_"
        )

        keyboard = {
            "inline_keyboard": [
                [
                    {"text": f"{self.EMOJI['accept']} Accept",               "callback_data": f"acc:{sha}"},
                    {"text": f"{self.EMOJI['decline']} Decline & Rollback",  "callback_data": f"dec:{sha}"},
                ],
                [
                    {"text": f"{self.EMOJI['report']} Transparency Report",  "callback_data": f"rep:{sha}"},
                ],
            ]
        }
        # Bug 3 fix: Telegram hard limit is 4096 chars. Truncate the body (not
        # the buttons) so the card always sends even on large commits.
        MAX_TELEGRAM = 4000  # leave headroom for markup overhead
        if len(text) > MAX_TELEGRAM:
            text = text[:MAX_TELEGRAM] + "\n\n_…message truncated_"

        return await self.send_message(chat_id, text, reply_markup=keyboard)

    # ── Button action handlers ─────────────────────────────────────────────────

    async def handle_accept(
        self, chat_id, callback_query_id, message_id,
        commit_metadata, decision, on_accept,
    ) -> None:
        sha = commit_metadata.get("sha", "N/A")[:7]
        await self.answer_callback(callback_query_id, f"✅ Commit {sha} accepted!")
        proc = await self.send_processing(chat_id, "✅ _Accepting commit…_")
        try:
            await on_accept()
        except Exception as exc:
            logger.error("Accept callback failed: %s", exc)
        finally:
            await self.delete_message(chat_id, proc)
        await self.edit_message(
            chat_id, message_id,
            f"✅ *COMMIT ACCEPTED*\n\n"
            f"`{sha}` — {commit_metadata.get('message','')[:80]}\n"
            f"by {commit_metadata.get('author_name','Unknown')}\n\n"
            f"{self._risk_emoji(decision.risk_level)} Risk: {decision.risk_level.upper()}\n"
            f"✅ Commit kept in repository\n\n"
            f"_{datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M:%S UTC')}_",
            reply_markup={"inline_keyboard": []},
        )

    async def handle_decline(
        self, chat_id, callback_query_id, message_id,
        commit_metadata, decision, on_decline,
    ) -> None:
        sha = commit_metadata.get("sha", "N/A")[:7]
        await self.answer_callback(callback_query_id, f"⏪ Rolling back {sha}…")
        proc = await self.send_processing(chat_id, f"⏪ _Rolling back `{sha}`…_")
        result: Dict[str, Any] = {}
        error_msg = None
        try:
            result = await on_decline()
        except Exception as exc:
            error_msg = str(exc)
        finally:
            await self.delete_message(chat_id, proc)

        if result.get("success"):
            text = (
                f"❌ *COMMIT DECLINED & ROLLED BACK*\n\n"
                f"`{sha}` — {commit_metadata.get('message','')[:80]}\n"
                f"by {commit_metadata.get('author_name','Unknown')}\n\n"
                f"{self._risk_emoji(decision.risk_level)} Risk: {decision.risk_level.upper()}\n"
                f"⏪ Strategy: {result.get('strategy','unknown')}\n✅ Rolled back\n"
            )
            if result.get("revert_sha"):
                text += f"📝 Revert SHA: `{result['revert_sha'][:7]}`\n"
        else:
            text = (
                f"❌ *DECLINED — ROLLBACK FAILED*\n\n"
                f"`{sha}` — {commit_metadata.get('message','')[:80]}\n\n"
                f"❌ Error: {error_msg or result.get('message','Unknown')}\n\n"
                f"⚠️ Please rollback manually.\n{commit_metadata.get('url','')}"
            )
        await self.edit_message(chat_id, message_id, text, reply_markup={"inline_keyboard": []})

    async def handle_report(
        self, chat_id, callback_query_id, commit_metadata, decision,
    ) -> None:
        sha = commit_metadata.get("sha", "N/A")[:7]
        await self.answer_callback(callback_query_id, "Generating report…")
        proc = await self.send_processing(chat_id, f"📊 _Generating report for `{sha}`…_")

        try:
            repo_full    = commit_metadata.get("_repo_full_name", "")
            if "/" not in repo_full:
                raise ValueError(f"Missing repo info in commit metadata: {repo_full!r}")
            owner, repo  = repo_full.split("/", 1)
            user         = db.get_user(chat_id)
            from github_service import GitHubService
            gh           = GitHubService(token=user["github_token"]) if user else None
            repo_context = await gh.fetch_repo_context(owner, repo) if gh else {}
            report_text  = await ai_service.generate_transparency_report(commit_metadata, repo_context, decision)
            docx_path    = await self._build_report_docx(commit_metadata, decision, report_text)
            await self.delete_message(chat_id, proc)
            await self.send_document(
                chat_id, docx_path,
                caption=f"📊 Transparency Report — `{sha}` | {decision.risk_level.upper()} risk",
                filename=f"report_{sha}.docx",
            )
            try:
                os.unlink(docx_path)
            except OSError:
                pass
            if gh:
                await gh.close()
        except Exception as exc:
            logger.error("Report failed: %s", exc)
            await self.delete_message(chat_id, proc)
            await self.send_message(
                chat_id,
                f"⚠️ *Report Error*\n\n`{self._safe_error(exc)}`\n\n"
                f"*Basic:* {decision.decision.upper()} | {decision.risk_level.upper()} | "
                f"{decision.confidence_score:.0%}\n\n{decision.summary}",
            )

    # ── Full codebase analysis ─────────────────────────────────────────────────

    async def handle_full_code_analysis(self, chat_id: str) -> None:
        """
        Triggered by the '🔍 Full Code Analysis' menu button.
        Fetches repo context, sends everything to AI in one compact call,
        then generates a structured .docx report and delivers it.
        """
        user = db.get_user(chat_id)
        if not user or not db.is_setup_complete(chat_id):
            await self.send_message(
                chat_id,
                "⚠️ Please complete setup first — send /start",
            )
            return

        owner = user.get("owner", "")
        repo  = user.get("repo", "")
        proc  = await self.send_processing(
            chat_id,
            f"🔍 _Starting full code analysis for `{owner}/{repo}`…_\n\n"
            "_This may take 30–60 seconds. Fetching codebase context…_",
        )

        try:
            from github_service import GitHubService
            gh = GitHubService(token=user["github_token"])

            await self.edit_message(chat_id, proc, "🔍 _Loading repository structure and key files…_")
            repo_context = await gh.fetch_repo_context(owner, repo, default_branch=user.get("branch", "main"))
            await gh.close()

            # All historical reviews for this user
            reviews = db.get_all_reviews_for_user(chat_id, limit=200)

            await self.edit_message(chat_id, proc, "🤖 _AI is auditing the entire codebase… (may take ~30s)_")
            analysis = await ai_service.analyze_full_codebase(repo_context, reviews)

            await self.edit_message(chat_id, proc, "📄 _Generating code analysis report…_")
            docx_path = await self._build_code_analysis_docx(owner, repo, analysis, reviews)

            await self.delete_message(chat_id, proc)
            health = analysis.get("overall_health", "unknown").upper()
            await self.send_document(
                chat_id, docx_path,
                caption=(
                    f"🔍 *Full Code Analysis — `{owner}/{repo}`*\n\n"
                    f"Overall Health: *{health}*\n"
                    f"Security Score: *{analysis.get('security', {}).get('score', '?')}/100*\n"
                    f"Quality Score: *{analysis.get('code_quality', {}).get('score', '?')}/100*\n\n"
                    f"_Full report in the document above._"
                ),
                filename=f"code_analysis_{repo}.docx",
            )
            try:
                os.unlink(docx_path)
            except OSError:
                pass

        except Exception as exc:
            logger.error("Full code analysis failed for %s: %s", chat_id, exc)
            await self.delete_message(chat_id, proc)
            await self.send_message(
                chat_id,
                f"⚠️ *Code Analysis Failed*\n\n`{self._safe_error(exc)}`\n\n"
                "_Make sure your GitHub token has repo read access._",
            )

    async def handle_author_review(self, chat_id: str) -> None:
        """
        Triggered by the '👥 Author Performance Review' menu button.
        Pulls all stored review records, aggregates per-author stats,
        asks AI for qualitative assessment, then sends a .docx report.
        No live GitHub API calls needed — uses the local DB only.
        """
        user = db.get_user(chat_id)
        if not user or not db.is_setup_complete(chat_id):
            await self.send_message(
                chat_id,
                "⚠️ Please complete setup first — send /start",
            )
            return

        owner = user.get("owner", "")
        repo  = user.get("repo", "")
        proc  = await self.send_processing(
            chat_id,
            f"👥 _Analysing author performance for `{owner}/{repo}`…_",
        )

        try:
            reviews = db.get_all_reviews_for_user(chat_id, limit=500)
            if not reviews:
                await self.delete_message(chat_id, proc)
                await self.send_message(
                    chat_id,
                    "👥 *Author Performance Review*\n\n"
                    "_No commit history found yet. Push some commits first!_",
                )
                return

            await self.edit_message(chat_id, proc, "🤖 _AI is evaluating each author's track record…_")
            analysis = await ai_service.analyze_authors(reviews, repo_name=f"{owner}/{repo}")

            await self.edit_message(chat_id, proc, "📄 _Generating author performance report…_")
            docx_path = await self._build_author_review_docx(owner, repo, analysis)

            await self.delete_message(chat_id, proc)
            n_authors = len(analysis.get("authors") or [])
            mvp = analysis.get("mvp") or "N/A"
            await self.send_document(
                chat_id, docx_path,
                caption=(
                    f"👥 *Author Performance Review — `{owner}/{repo}`*\n\n"
                    f"Authors analysed: *{n_authors}*\n"
                    f"🏆 MVP: *{mvp}*\n\n"
                    f"_Full breakdown in the document above._"
                ),
                filename=f"author_review_{repo}.docx",
            )
            try:
                os.unlink(docx_path)
            except OSError:
                pass

        except Exception as exc:
            logger.error("Author review failed for %s: %s", chat_id, exc)
            await self.delete_message(chat_id, proc)
            await self.send_message(
                chat_id,
                f"⚠️ *Author Review Failed*\n\n`{self._safe_error(exc)}`",
            )

    # ── Code analysis docx builder ────────────────────────────────────────────

    async def _build_code_analysis_docx(
        self, owner: str, repo: str,
        analysis: dict, reviews: list,
    ) -> str:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(
            None,
            lambda: self._build_code_analysis_docx_sync(owner, repo, analysis, reviews),
        )

    def _build_code_analysis_docx_sync(
        self, owner: str, repo: str, analysis: dict, reviews: list,
    ) -> str:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

        BLUE      = RGBColor(0x1F, 0x49, 0x7D)
        MID_BLUE  = RGBColor(0x2E, 0x75, 0xB6)
        GREEN     = RGBColor(0x37, 0x86, 0x10)
        RED       = RGBColor(0xC0, 0x00, 0x00)
        ORANGE    = RGBColor(0xC5, 0x5A, 0x11)
        YELLOW    = RGBColor(0xBF, 0x8F, 0x00)
        GREY      = RGBColor(0x88, 0x88, 0x88)
        BLACK     = RGBColor(0x00, 0x00, 0x00)

        HEALTH_COLOUR = {
            "excellent": GREEN, "good": GREEN, "fair": YELLOW,
            "poor": ORANGE, "critical": RED,
        }

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        def h1(text):
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = BLUE
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(16)

        def h2(text):
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = MID_BLUE
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(13)

        def kv(key, value, vc=BLACK):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            rk = p.add_run(f"{key}: ")
            rk.bold = True; rk.font.name = "Arial"; rk.font.size = Pt(11)
            rv = p.add_run(value)
            rv.font.name = "Arial"; rv.font.size = Pt(11)
            rv.font.color.rgb = vc

        def bullet(text):
            p = doc.add_paragraph(text, style="List Bullet")
            if p.runs:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(11)

        def body(text):
            p = doc.add_paragraph(text)
            if p.runs:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)

        def spacer():
            doc.add_paragraph()

        def score_bar(label, score):
            """Render a score as 'label: XX/100 ██████░░░░' in a single paragraph."""
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            rk = p.add_run(f"{label}: ")
            rk.bold = True; rk.font.name = "Arial"; rk.font.size = Pt(11)
            pct = max(0, min(100, int(score or 0)))
            filled = round(pct / 10)
            bar = "█" * filled + "░" * (10 - filled)
            rv = p.add_run(f"{pct}/100  {bar}")
            rv.font.name = "Arial"; rv.font.size = Pt(11)
            color = GREEN if pct >= 70 else (YELLOW if pct >= 40 else RED)
            rv.font.color.rgb = color

        # ── Title ─────────────────────────────────────────────────────────────
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run("Commit Guardian — Full Code Analysis Report")
        tr.bold = True; tr.font.name = "Arial"; tr.font.size = Pt(20)
        tr.font.color.rgb = BLUE

        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(f"{owner}/{repo}  ·  Generated {generated}")
        sr.italic = True; sr.font.name = "Arial"; sr.font.size = Pt(9)
        sr.font.color.rgb = GREY
        spacer()

        # ── Executive Summary ─────────────────────────────────────────────────
        h1("Executive Summary")
        health = analysis.get("overall_health", "unknown")
        kv("Overall Health", health.upper(), vc=HEALTH_COLOUR.get(health, BLACK))
        body(analysis.get("executive_summary", "No summary generated."))
        spacer()

        # ── Score Overview ─────────────────────────────────────────────────────
        h1("Score Overview")
        sec_score  = analysis.get("security", {}).get("score", 0)
        qual_score = analysis.get("code_quality", {}).get("score", 0)
        score_bar("Security Score", sec_score)
        score_bar("Code Quality Score", qual_score)
        spacer()

        # ── Security ─────────────────────────────────────────────────────────
        h1("Security Assessment")
        sec = analysis.get("security", {})
        body(sec.get("summary", ""))
        if sec.get("findings"):
            h2("Findings")
            for f in sec["findings"]:
                bullet(f)
        if sec.get("recommendations"):
            h2("Security Recommendations")
            for r in sec["recommendations"]:
                bullet(r)
        spacer()

        # ── Code Quality ──────────────────────────────────────────────────────
        h1("Code Quality")
        cq = analysis.get("code_quality", {})
        body(cq.get("summary", ""))
        if cq.get("strengths"):
            h2("Strengths")
            for s in cq["strengths"]:
                bullet(s)
        if cq.get("weaknesses"):
            h2("Areas for Improvement")
            for w in cq["weaknesses"]:
                bullet(w)
        spacer()

        # ── Architecture ──────────────────────────────────────────────────────
        h1("Architecture")
        arch = analysis.get("architecture", {})
        body(arch.get("summary", ""))
        if arch.get("patterns_detected"):
            h2("Patterns Detected")
            for p in arch["patterns_detected"]:
                bullet(p)
        if arch.get("concerns"):
            h2("Architectural Concerns")
            for c in arch["concerns"]:
                bullet(c)
        spacer()

        # ── Progress / Commit History ─────────────────────────────────────────
        h1("Progress & Review History")
        prog = analysis.get("progress", {})
        body(prog.get("summary", ""))
        total   = prog.get("total_commits_reviewed", len(reviews))
        accepted = prog.get("accepted", sum(1 for r in reviews if r.get("status") == "accepted"))
        declined = prog.get("declined", sum(1 for r in reviews if r.get("status") == "declined"))
        pending  = prog.get("pending", sum(1 for r in reviews if r.get("status") == "pending"))
        kv("Total Commits Reviewed", str(total))
        kv("Accepted",  str(accepted), vc=GREEN)
        kv("Declined",  str(declined), vc=RED)
        kv("Pending",   str(pending),  vc=YELLOW)
        kv("High Risk", str(prog.get("high_risk_commits", "?")))
        kv("Trend",     prog.get("trend", "insufficient_data").replace("_", " ").title())
        spacer()

        # ── Dependencies ──────────────────────────────────────────────────────
        h1("Dependencies")
        deps = analysis.get("dependencies", {})
        body(deps.get("summary", ""))
        if deps.get("notable"):
            for n in deps["notable"]:
                bullet(n)
        spacer()

        # ── Top Recommendations ───────────────────────────────────────────────
        h1("Top Recommendations")
        for rec in (analysis.get("top_recommendations") or []):
            bullet(rec)

        fd, out_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(out_path)
        return out_path

    # ── Author performance docx builder ───────────────────────────────────────

    async def _build_author_review_docx(
        self, owner: str, repo: str, analysis: dict,
    ) -> str:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(
            None,
            lambda: self._build_author_review_docx_sync(owner, repo, analysis),
        )

    def _build_author_review_docx_sync(
        self, owner: str, repo: str, analysis: dict,
    ) -> str:
        generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

        BLUE     = RGBColor(0x1F, 0x49, 0x7D)
        MID_BLUE = RGBColor(0x2E, 0x75, 0xB6)
        GREEN    = RGBColor(0x37, 0x86, 0x10)
        RED      = RGBColor(0xC0, 0x00, 0x00)
        ORANGE   = RGBColor(0xC5, 0x5A, 0x11)
        YELLOW   = RGBColor(0xBF, 0x8F, 0x00)
        GREY     = RGBColor(0x88, 0x88, 0x88)
        BLACK    = RGBColor(0x00, 0x00, 0x00)
        GOLD     = RGBColor(0xFF, 0xD7, 0x00)

        RATING_COLOUR = {
            "excellent":     GREEN,
            "good":          GREEN,
            "average":       YELLOW,
            "below_average": ORANGE,
            "poor":          RED,
        }

        doc = Document()
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        def h1(text):
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = BLUE
            p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(16)

        def h2(text):
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = MID_BLUE
            p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(13)

        def kv(key, value, vc=BLACK):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            rk = p.add_run(f"{key}: ")
            rk.bold = True; rk.font.name = "Arial"; rk.font.size = Pt(11)
            rv = p.add_run(value)
            rv.font.name = "Arial"; rv.font.size = Pt(11)
            rv.font.color.rgb = vc

        def bullet(text):
            p = doc.add_paragraph(text, style="List Bullet")
            if p.runs:
                p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(11)

        def body(text):
            p = doc.add_paragraph(text)
            if p.runs:
                p.runs[0].font.name = "Arial"; p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)

        def spacer():
            doc.add_paragraph()

        # ── Title ─────────────────────────────────────────────────────────────
        tp = doc.add_paragraph()
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tr = tp.add_run("Commit Guardian — Author Performance Review")
        tr.bold = True; tr.font.name = "Arial"; tr.font.size = Pt(20)
        tr.font.color.rgb = BLUE

        sp = doc.add_paragraph()
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sp.add_run(f"{owner}/{repo}  ·  Generated {generated}")
        sr.italic = True; sr.font.name = "Arial"; sr.font.size = Pt(9)
        sr.font.color.rgb = GREY
        spacer()

        # ── Team Summary ──────────────────────────────────────────────────────
        h1("Team Overview")
        body(analysis.get("team_summary", "No summary generated."))

        mvp = analysis.get("mvp")
        if mvp:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            rk = p.add_run("🏆 MVP: ")
            rk.bold = True; rk.font.name = "Arial"; rk.font.size = Pt(12)
            rv = p.add_run(mvp)
            rv.bold = True; rv.font.name = "Arial"; rv.font.size = Pt(12)
            rv.font.color.rgb = GOLD

        needs_attn = analysis.get("needs_attention")
        if needs_attn:
            kv("⚠️  Needs Attention", needs_attn, vc=ORANGE)
        spacer()

        # ── Per-author sections ───────────────────────────────────────────────
        h1("Individual Author Reports")

        authors = analysis.get("authors") or []
        # Sort: best performers first
        rating_order = {"excellent": 0, "good": 1, "average": 2, "below_average": 3, "poor": 4}
        authors_sorted = sorted(
            authors,
            key=lambda a: rating_order.get(a.get("performance_rating", "average"), 2),
        )

        for author in authors_sorted:
            name   = author.get("name", "Unknown")
            rating = author.get("performance_rating", "average")
            colour = RATING_COLOUR.get(rating, BLACK)

            # Author heading with rating badge
            hp = doc.add_heading(level=2)
            hp.paragraph_format.space_before = Pt(12)
            hr = hp.runs[0] if hp.runs else hp.add_run()
            hr.text = name
            hr.font.color.rgb = MID_BLUE
            hr.font.name = "Arial"
            hr.font.size = Pt(13)
            badge_run = hp.add_run(f"  [{rating.replace('_', ' ').upper()}]")
            badge_run.font.color.rgb = colour
            badge_run.font.name = "Arial"
            badge_run.font.size = Pt(11)
            badge_run.bold = True

            total   = author.get("total_commits", 0)
            accepted = author.get("accepted", 0)
            declined = author.get("declined", 0)
            pending  = author.get("pending", 0)
            high_risk = author.get("high_risk_commits", 0)
            dr      = author.get("decline_rate_pct", 0)

            kv("Total Commits",   str(total))
            kv("Accepted",        str(accepted), vc=GREEN)
            kv("Declined",        str(declined), vc=RED if declined > 0 else BLACK)
            kv("Pending",         str(pending),  vc=YELLOW if pending > 0 else BLACK)
            kv("High-Risk Commits", str(high_risk), vc=ORANGE if high_risk > 0 else BLACK)
            kv("Decline Rate",    f"{dr}%", vc=RED if dr >= 40 else (ORANGE if dr >= 20 else GREEN))

            if author.get("verdict"):
                body(author["verdict"])

            if author.get("strengths"):
                p = doc.add_paragraph()
                p.add_run("Strengths:").bold = True
                for s in author["strengths"]:
                    bullet(s)

            if author.get("concerns"):
                p = doc.add_paragraph()
                p.add_run("Concerns:").bold = True
                for c in author["concerns"]:
                    bullet(c)

            spacer()

        fd, out_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(out_path)
        return out_path

    # ── Word report builder ───────────────────────────────────────────────────

    async def _build_report_docx(self, commit_metadata, decision, report_text) -> str:
        """
        Build a .docx transparency report using python-docx (pure Python, no Node.js).
        Runs in a thread executor so it doesn't block the event loop.
        """
        loop = asyncio.get_running_loop()
        out_path = await loop.run_in_executor(
            None,
            lambda: self._build_report_docx_sync(commit_metadata, decision, report_text),
        )
        return out_path

    def _build_report_docx_sync(self, commit_metadata, decision, report_text) -> str:
        """Synchronous docx builder — called from a thread executor."""
        sha       = commit_metadata.get("sha", "unknown")[:7]
        author    = commit_metadata.get("author_name", "Unknown")
        email     = commit_metadata.get("author_email", "N/A")
        committed = commit_metadata.get("committed_at", "N/A")
        msg       = commit_metadata.get("message", "")[:300]
        files     = commit_metadata.get("files", [])[:20]
        generated = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S UTC")

        # ── Colour palette ────────────────────────────────────────────────────
        BLUE      = RGBColor(0x1F, 0x49, 0x7D)   # heading 1
        MID_BLUE  = RGBColor(0x2E, 0x75, 0xB6)   # heading 2
        GREY      = RGBColor(0x88, 0x88, 0x88)   # subtitle
        BLACK     = RGBColor(0x00, 0x00, 0x00)

        RISK_COLOUR = {
            "critical": RGBColor(0xC0, 0x00, 0x00),
            "high":     RGBColor(0xC5, 0x5A, 0x11),
            "medium":   RGBColor(0xBF, 0x8F, 0x00),
            "low":      RGBColor(0x37, 0x86, 0x10),
        }

        doc = Document()

        # ── Page margins (1 inch all round) ───────────────────────────────────
        for section in doc.sections:
            section.top_margin    = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin   = Inches(1)
            section.right_margin  = Inches(1)

        # ── Default body style ────────────────────────────────────────────────
        style = doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(11)

        # ── Helpers ───────────────────────────────────────────────────────────
        def h1(text: str):
            p = doc.add_heading(text, level=1)
            p.runs[0].font.color.rgb = BLUE
            p.runs[0].font.name      = "Arial"
            p.runs[0].font.size      = Pt(16)

        def h2(text: str):
            p = doc.add_heading(text, level=2)
            p.runs[0].font.color.rgb = MID_BLUE
            p.runs[0].font.name      = "Arial"
            p.runs[0].font.size      = Pt(13)

        def kv(key: str, value: str, value_colour: RGBColor = BLACK):
            """Bold key followed by plain value on the same line."""
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run_k = p.add_run(f"{key}: ")
            run_k.bold           = True
            run_k.font.name      = "Arial"
            run_k.font.size      = Pt(11)
            run_v = p.add_run(value)
            run_v.font.name      = "Arial"
            run_v.font.size      = Pt(11)
            run_v.font.color.rgb = value_colour

        def bullet(text: str):
            p = doc.add_paragraph(text, style="List Bullet")
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(11)

        def body(text: str):
            p = doc.add_paragraph(text)
            if p.runs:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)

        def spacer():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)

        # ── Title block ───────────────────────────────────────────────────────
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_p.add_run("Commit Guardian — Transparency Report")
        title_run.bold           = True
        title_run.font.name      = "Arial"
        title_run.font.size      = Pt(20)
        title_run.font.color.rgb = BLUE

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(f"Generated {generated}")
        sub_run.italic           = True
        sub_run.font.name        = "Arial"
        sub_run.font.size        = Pt(9)
        sub_run.font.color.rgb   = GREY
        spacer()

        # ── Commit Details ────────────────────────────────────────────────────
        h1("Commit Details")
        kv("SHA",       sha)
        kv("Author",    f"{author} <{email}>")
        kv("Committed", committed)
        kv("Message",   msg)
        spacer()

        # ── AI Assessment ─────────────────────────────────────────────────────
        h1("AI Assessment")
        risk_col = RISK_COLOUR.get(decision.risk_level.lower(), BLACK)
        kv("Decision",   decision.decision.upper())
        kv("Risk Level", decision.risk_level.upper(), value_colour=risk_col)
        kv("Confidence", f"{decision.confidence_score:.0%}")
        kv("Summary",    decision.summary)
        spacer()

        # ── Concerns ──────────────────────────────────────────────────────────
        if decision.concerns:
            h2("Concerns")
            for c in decision.concerns[:10]:
                bullet(c)
            spacer()

        # ── Positive Aspects ──────────────────────────────────────────────────
        if decision.positive_aspects:
            h2("Positive Aspects")
            for p_item in decision.positive_aspects[:10]:
                bullet(p_item)
            spacer()

        # ── Recommendations ───────────────────────────────────────────────────
        if decision.recommendations:
            h2("Recommendations")
            for r in decision.recommendations[:10]:
                bullet(r)
            spacer()

        # ── Changed Files ─────────────────────────────────────────────────────
        h1("Changed Files")
        for f in files:
            fname   = f.get("filename", "unknown")
            status  = f.get("status", "")
            adds    = f.get("additions", 0)
            dels    = f.get("deletions", 0)
            body(f"{fname}  |  {status}  |  +{adds} / -{dels}")
        spacer()

        # ── Detailed Analysis ─────────────────────────────────────────────────
        h1("Detailed Analysis")
        for line in report_text.split("\n"):
            if line.strip():
                body(line)

        # ── Save ──────────────────────────────────────────────────────────────
        import contextlib
        fd, out_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        try:
            doc.save(out_path)
        except Exception:
            with contextlib.suppress(OSError):
                os.unlink(out_path)
            raise  # re-raise so _build_report_docx / handle_report catches it
        return out_path


telegram_service = TelegramService()
